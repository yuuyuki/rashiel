import argparse
import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Callable

ENCODING = 'utf-8'
CONFIG_FILE = 'config.json'
SCRIPT_FILE_TYPE = '.utf'

# Ensure custom_function has function type that takes a string and returns a string
custom_function: Callable[[str], str] = None

def combine_script_to_docx(output_filename):
    doc = Document()
    for file in Path('.').glob(f"*{SCRIPT_FILE_TYPE}"):
        doc.add_paragraph(file.name).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph('-' * 40)

        with file.open('r', encoding=ENCODING, errors='ignore') as f:
            for line in f:
                doc.add_paragraph(line)

    doc.save(output_filename)
    print(f"Combined files into {output_filename}")

def split_docx_to_scripts(input_filename):
    doc = Document(input_filename)
    script_files = {}
    current_filename = None
    content_lines = []

    # Parse .docx into a mapping of filename -> list of lines
    for para in doc.paragraphs:
        if para.text.endswith(SCRIPT_FILE_TYPE) and len(para.text.strip()) > 4:
            if current_filename and content_lines:  # add prev file content
                script_files[current_filename] = content_lines
            current_filename = para.text.strip()
            content_lines = []
        elif para.text == '-' * 40:
            continue
        else:
            content_lines.append(para.text)

    if current_filename and content_lines:   # add last file content
        script_files[current_filename] = content_lines

    for filename, docx_lines in script_files.items():
        with open(filename, 'w', encoding=ENCODING, errors='ignore') as f:
            f.writelines(docx_lines)

def process_script_files():
    """Process all script files in the current directory with custom_function, overwriting them."""
    if custom_function is None:
        print("Error: custom_function is not defined.")
        return
    
    for file in Path('.').glob('*'+ SCRIPT_FILE_TYPE):
        with file.open('r', encoding=ENCODING, errors='ignore') as f:
            lines = f.readlines()

        processed_lines = []
        for line in lines:
            stripped = line.strip()
            if stripped == '' or stripped.startswith(';') or stripped.startswith('#'):
                processed_lines.append(line)
            else:
                processed_line = custom_function(line.rstrip('\n')) + '\n'
                processed_lines.append(processed_line)

        # Warning if line counts differ
        if len(processed_lines) != len(lines):
            print(f"Warning: Line count mismatch for {file.name}. Original: {len(lines)}, Processed: {len(processed_lines)}")

        with file.open('w', encoding=ENCODING, errors='ignore') as f:
            f.writelines(processed_lines)

        print(f"Overwritten: {file.name}")


def main():
    parser = argparse.ArgumentParser(description="Combine or split script files and .docx documents.")
    subparsers = parser.add_subparsers(dest='command')

    combine_parser = subparsers.add_parser('combine', help='Combine all script files into a .docx')
    combine_parser.add_argument('output', help='Output .docx filename')

    split_parser = subparsers.add_parser('split', help='Split .docx into script files with line replacements')
    split_parser.add_argument('input', help='Input .docx filename')

    process_parser = subparsers.add_parser('process', help='Process script files in-place using custom_function on non-comment, non-blank lines')

    args = parser.parse_args()

    if args.command == 'combine':
        combine_script_to_docx(args.output)
    elif args.command == 'split':
        split_docx_to_scripts(args.input)
    elif args.command == 'process':
        process_script_files()
    else:
        parser.print_help()

if __name__ == '__main__':
    main()
